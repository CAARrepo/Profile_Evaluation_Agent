"""Extract plain text from uploaded PDF / DOCX resumes and evidence."""

from __future__ import annotations

from pathlib import Path


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(path)
    if suffix in {".docx", ".doc"}:
        return _extract_docx(path)
    if suffix in {".txt", ".md", ".csv"}:
        return path.read_text(encoding="utf-8", errors="ignore")
    # Some resumes are stored as `.docx.pdf` style names; try PDF first.
    if ".pdf" in path.name.lower():
        try:
            return _extract_pdf(path)
        except Exception:
            pass
    raise ValueError(f"Unsupported document type: {path.name}")


def _extract_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    parts: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            parts.append(text)
    return "\n".join(parts).strip()


def _extract_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    # Also pull simple table text
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()
